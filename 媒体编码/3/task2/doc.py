from docx import Document
from docx.shared import Inches

# 创建 Word 文档

doc = Document()
doc.add_heading('JPEG 量化步长优化分析报告', 0)

# 1. 实验结果

doc.add_heading('1. 实验结果', level=1)
doc.add_paragraph('对 Lena 图像，使用标准亮度量化表及步长调节后的压缩效果：')
doc.add_paragraph(
'标准表: PSNR=35.75 dB, 码率=0.108\n'
'步长缩小: PSNR=37.77 dB, 码率=0.169\n'
'步长放大: PSNR=33.67 dB, 码率=0.070'
)

# 2. 重建图像

doc.add_heading('2. 重建图像对比', level=1)
doc.add_paragraph('下图展示了不同量化表重建的 Lena 图像，步长放大量化表可见明显块效应。')
doc.add_picture('recon_half.png', width=Inches(3))
doc.add_picture('recon_std.png', width=Inches(3))
doc.add_picture('recon_double.png', width=Inches(3))

# 3. 码率 - PSNR 折线图

doc.add_heading('3. 码率 - PSNR 折线图', level=1)
doc.add_picture('rate_psnr.png', width=Inches(6))

# 4. 分析总结

doc.add_heading('4. 分析总结', level=1)
doc.add_paragraph(
'实验结果显示：\n'
'- 量化步长缩小: 码率升高，PSNR 提高，图像细节保留更多\n'
'- 量化步长放大: 码率降低 ≥30%，PSNR 降低约 2dB，出现块效应\n'
'- 量化步长越大，更多 DCT 系数被量化为零，压缩比增大，但重建质量下降\n'
'- 量化步长越小，保留更多细节，但码率上升'
)

# 保存文档

doc.save('JPEG_Quant_Optimization_Report.docx')
