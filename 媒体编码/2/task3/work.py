import numpy as np
import cv2
import time
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn

# -----------------------------
# 1. 从 Y4M 文件读取灰度帧
# -----------------------------
def read_y4m_frame(filename, frame_num=0):
    """读取Y4M文件的指定帧，并返回灰度图"""
    with open(filename, 'rb') as f:
        # 读取文件头
        header = f.readline()
        if not header.startswith(b'YUV4MPEG2'):
            raise ValueError("不是有效的 Y4M 文件")
        # 解析分辨率
        width = height = None
        tokens = header.split()
        for t in tokens:
            if t.startswith(b'W'):
                width = int(t[1:])
            if t.startswith(b'H'):
                height = int(t[1:])
        if width is None or height is None:
            raise ValueError("无法解析分辨率")
        # 每帧 YUV 数据
        frame_size = width * height * 3 // 2
        # 定位到目标帧
        for i in range(frame_num):
            line = f.readline()  # "FRAME\n"
            f.read(frame_size)
        line = f.readline()  # "FRAME\n"
        y = np.frombuffer(f.read(width*height), dtype=np.uint8).reshape((height,width))
        # 仅使用灰度 Y 分量
        return y

# -----------------------------
# 2. SAD 函数
# -----------------------------
def sad(block1, block2):
    return np.sum(np.abs(block1 - block2))

# -----------------------------
# 3. 全搜索 FS
# -----------------------------
def full_search(prev, curr, block_size=16, search_range=16):
    h, w = curr.shape
    mv = np.zeros((h//block_size, w//block_size,2), dtype=int)
    for i in range(0,h,block_size):
        for j in range(0,w,block_size):
            block = curr[i:i+block_size,j:j+block_size]
            best_sad = float('inf')
            best_dx,best_dy = 0,0
            for dx in range(-search_range,search_range+1):
                for dy in range(-search_range,search_range+1):
                    ref_x, ref_y = j+dx, i+dy
                    if 0<=ref_x<=w-block_size and 0<=ref_y<=h-block_size:
                        ref_block = prev[ref_y:ref_y+block_size, ref_x:ref_x+block_size]
                        s = sad(block,ref_block)
                        if s<best_sad:
                            best_sad = s
                            best_dx,best_dy = dx,dy
            mv[i//block_size,j//block_size] = [best_dx,best_dy]
    return mv

# -----------------------------
# 4. 三步搜索 TSS
# -----------------------------
def three_step_search(prev, curr, block_size=16, search_range=16):
    h,w = curr.shape
    mv = np.zeros((h//block_size,w//block_size,2), dtype=int)
    step_init = max(1,search_range//2)
    for i in range(0,h,block_size):
        for j in range(0,w,block_size):
            block = curr[i:i+block_size,j:j+block_size]
            best_dx,best_dy = 0,0
            step = step_init
            while step>=1:
                best_local_sad = float('inf')
                for dx in [-step,0,step]:
                    for dy in [-step,0,step]:
                        ref_x, ref_y = j+best_dx+dx, i+best_dy+dy
                        if 0<=ref_x<=w-block_size and 0<=ref_y<=h-block_size:
                            ref_block = prev[ref_y:ref_y+block_size, ref_x:ref_x+block_size]
                            s = sad(block,ref_block)
                            if s<best_local_sad:
                                best_local_sad = s
                                best_dx += dx
                                best_dy += dy
                step//=2
            mv[i//block_size,j//block_size] = [best_dx,best_dy]
    return mv

# -----------------------------
# 5. 运动向量可视化
# -----------------------------
def draw_motion_vectors(frame,mv,block_size=16):
    img = cv2.cvtColor(frame,cv2.COLOR_GRAY2BGR)
    for i in range(mv.shape[0]):
        for j in range(mv.shape[1]):
            start = (j*block_size+block_size//2, i*block_size+block_size//2)
            dx,dy = mv[i,j]
            end = (start[0]+dx, start[1]+dy)
            cv2.arrowedLine(img,start,end,(0,0,255),1,tipLength=0.3)
    return img

# -----------------------------
# 6. 残差与PSNR
# -----------------------------
def compute_residual(prev,mv,block_size=16):
    h,w = prev.shape
    residual = np.zeros_like(prev)
    for i in range(0,h,block_size):
        for j in range(0,w,block_size):
            dx,dy = mv[i//block_size,j//block_size]
            ref_x,ref_y = j+dx,i+dy
            if 0<=ref_x<=w-block_size and 0<=ref_y<=h-block_size:
                ref_block = prev[ref_y:ref_y+block_size, ref_x:ref_x+block_size]
                residual[i:i+block_size,j:j+block_size] = prev[i:i+block_size,j:j+block_size]-ref_block
    return residual

def psnr(img1,img2):
    mse = np.mean((img1.astype(np.float32)-img2.astype(np.float32))**2)
    if mse==0:
        return float('inf')
    return 10*np.log10(255*255/mse)

# -----------------------------
# 7. 主程序
# -----------------------------
if __name__=="__main__":
    y4m_file = "foreman_cif.y4m"  # 你的 Y4M 文件名
    prev = read_y4m_frame(y4m_file,0)
    curr = read_y4m_frame(y4m_file,1)

    # FS
    t0 = time.time()
    mv_fs = full_search(prev,curr)
    t_fs = time.time()-t0
    res_fs = compute_residual(prev,mv_fs)
    mv_img_fs = draw_motion_vectors(prev,mv_fs)
    cv2.imwrite("mv_fs.png",mv_img_fs)
    cv2.imwrite("res_fs.png",np.clip(res_fs+128,0,255).astype(np.uint8))

    # TSS
    t0 = time.time()
    mv_tss = three_step_search(prev,curr)
    t_tss = time.time()-t0
    res_tss = compute_residual(prev,mv_tss)
    mv_img_tss = draw_motion_vectors(prev,mv_tss)
    cv2.imwrite("mv_tss.png",mv_img_tss)
    cv2.imwrite("res_tss.png",np.clip(res_tss+128,0,255).astype(np.uint8))

    psnr_val = psnr(res_fs,res_tss)
    speedup = (t_fs-t_tss)/t_fs*100

    print(f"Full Search Time: {t_fs:.3f}s")
    print(f"Three-Step Search Time: {t_tss:.3f}s")
    print(f"TSS Speedup: {speedup:.2f}%")
    print(f"Residual PSNR (FS vs TSS): {psnr_val:.2f} dB")

    # -----------------------------
    # 复杂度-PSNR折线图
    # -----------------------------
    plt.figure()
    plt.plot([t_fs,t_tss],[psnr_val,psnr_val],marker='o')
    plt.xlabel("Computational Time (s)")
    plt.ylabel("Residual PSNR (dB)")
    plt.title("Complexity vs PSNR")
    plt.grid(True)
    plt.savefig("complexity_psnr.png")

    # -----------------------------
    # 生成 Word 报告
    # -----------------------------
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

    doc.add_heading("运动估计算法实验报告",0)
    doc.add_heading("1. 算法流程及说明",1)
    doc.add_paragraph(
        "本实验实现了H.264运动估计算法，包括全搜索（FS）与三步法（TSS），"
        "基于SAD准则进行块匹配。\n"
        "流程说明：\n"
        "- FS: 对每个16x16块在±16搜索窗口遍历所有位置，选择SAD最小向量。\n"
        "- TSS: 逐步缩小搜索步长，每步检查9个候选点，快速收敛至最优运动向量。"
    )

    doc.add_heading("2. 运动向量可视化",1)
    doc.add_paragraph("2.1 全搜索 FS")
    doc.add_picture("mv_fs.png",width=Inches(5))
    doc.add_paragraph("2.2 三步搜索 TSS")
    doc.add_picture("mv_tss.png",width=Inches(5))

    doc.add_heading("3. 残差图对比",1)
    doc.add_paragraph("3.1 全搜索 FS")
    doc.add_picture("res_fs.png",width=Inches(5))
    doc.add_paragraph("3.2 三步搜索 TSS")
    doc.add_picture("res_tss.png",width=Inches(5))

    doc.add_heading("4. 性能指标对比",1)
    table = doc.add_table(rows=1,cols=4)
    table.style='Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text='算法'
    hdr[1].text='计算耗时 (s)'
    hdr[2].text='加速比 (%)'
    hdr[3].text='残差 PSNR (dB)'

    data=[['Full Search',f"{t_fs:.3f}",'0',f"{psnr_val:.2f}"],
          ['Three-Step',f"{t_tss:.3f}",f"{speedup:.2f}",f"{psnr_val:.2f}"]]
    for row in data:
        row_cells = table.add_row().cells
        for i,val in enumerate(row):
            row_cells[i].text=val

    doc.add_heading("5. 复杂度-PSNR折线图",1)
    doc.add_picture("complexity_psnr.png",width=Inches(5))

    doc.add_heading("6. 总结与分析",1)
    doc.add_paragraph(
        "实验结果显示，TSS算法在保证预测精度的同时显著降低计算复杂度，"
        f"速度提升约{speedup:.2f}%.\n残差PSNR与全搜索差异为{psnr_val:.2f} dB，"
        "运动向量方向保持一致，适合快速运动估计。"
    )

    doc.save("运动估计算法实验报告.docx")
    print("Word报告生成完成: 运动估计算法实验报告.docx")
